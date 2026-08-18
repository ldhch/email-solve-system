"""Build LBORA importable assets from the source docx files.

Reads:
  LBORA_客服标准FAQ_V1.0.docx      -> Q1-Q24 standard QA pairs (internal
                                      chargeback entries Q25/Q26 excluded)
  LBORA_7月7日起客户沟通案例_V1.0.docx -> real English replies; 5 non-duplicate
                                      scenarios become extra QA pairs, and all
                                      10 replies are embedded in the KB doc as
                                      worked examples
  (curated) SOP V2.0 content        -> customer-safe knowledge base V2.0
                                      (internal back-office/KPI/management
                                      sections stripped)

Writes:
  qa_pairs_import.json                        (29 pairs: 24 FAQ + 5 cases)
  LBORA_客服知识库_纯净版_V2.0.docx

Run from the repo root:
  python3 backend/scripts/build_lbora_assets.py
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[2]

# FAQ Q number -> admin category (Q25/Q26 are internal-only, intentionally absent).
CATEGORY_BY_Q = {
    range(1, 5): "售前与政策",
    range(5, 8): "尺码与适配",
    range(8, 13): "退货与退款",
    range(13, 17): "质量与商品状态",
    range(17, 20): "发错货与换货",
    range(20, 24): "物流",
    range(24, 25): "物流保险",
}

# Case study -> extra QA pairs (only scenarios NOT already covered by the FAQ).
# `case_num` maps to the parsed English reply; answer comes from the docx.
CASE_QA = [
    (1, "Can you change my order to a different size", "售前与政策"),
    (2, "My delivery failed and I don't want to contact the carrier again", "物流"),
    (3, "I will file a dispute or chargeback with my bank", "拒付风险"),
    (4, "bait and switch, I will file a consumer complaint and a credit card dispute", "拒付风险"),
    (10, "The hat is severely damaged and I am traveling soon, a replacement will not work", "质量与商品状态"),
]

_Q_RE = re.compile(r"^Q(\d+)\.\s+(.*)$")
_SECTION_RE = re.compile(r"^[一二三四五六七八九十]+、(.+)$")
_CASE_RE = re.compile(r"^案例(\d+)｜(.+)$")


def _category_for(q_num: int) -> str | None:
    for rng, cat in CATEGORY_BY_Q.items():
        if q_num in rng:
            return cat
    return None  # Q25/Q26 and beyond: internal, skipped.


def parse_faq(doc_path: Path) -> list[dict]:
    """Extract Q1-Q24 as {question, answer, category} in English."""
    doc = docx.Document(str(doc_path))
    entries: list[dict] = []
    pending: dict | None = None

    def finalize() -> None:
        nonlocal pending
        if pending is not None:
            answers = [a for a in pending["answers"] if a]
            if answers:
                cat = _category_for(pending["num"])
                if cat is not None:  # skip internal Q25/Q26
                    entries.append(
                        {
                            "question": pending["q"],
                            "answer": "\n\n".join(answers).strip(),
                            "category": cat,
                        }
                    )
            pending = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _SECTION_RE.match(text):
            finalize()
            continue
        m = _Q_RE.match(text)
        if m:
            finalize()
            q_num = int(m.group(1))
            # English question = everything before the last " / " (Chinese part).
            parts = m.group(2).split(" / ")
            english = " / ".join(parts[:-1]).strip() if len(parts) > 1 else m.group(2).strip()
            pending = {"num": q_num, "q": english, "answers": []}
            continue
        if pending is not None:
            pending["answers"].append(text)
    finalize()
    return entries


def parse_cases(doc_path: Path) -> list[dict]:
    """Extract (num, title, english reply) from the case-study docx."""
    doc = docx.Document(str(doc_path))
    paras = [p.text.strip() for p in doc.paragraphs]

    # Split paragraphs into case blocks, stopping at section headings.
    blocks: list[dict] = []
    cur: dict | None = None
    for t in paras:
        m = _CASE_RE.match(t)
        if m:
            if cur:
                blocks.append(cur)
            cur = {"num": int(m.group(1)), "title": m.group(2), "lines": []}
        elif cur is not None:
            if _SECTION_RE.match(t):
                blocks.append(cur)
                cur = None
            else:
                cur["lines"].append(t)
    if cur is not None:
        blocks.append(cur)

    cases: list[dict] = []
    for b in blocks:
        try:
            idx = b["lines"].index("English Reply")
        except ValueError:
            continue
        raw = b["lines"][idx + 1:]
        while raw and not raw[0].strip():
            raw.pop(0)
        while raw and not raw[-1].strip():
            raw.pop()
        reply = "\n".join(raw).strip()
        if reply:
            cases.append({"num": b["num"], "title": b["title"], "reply": reply})
    return cases


# ---------------------------------------------------------------- clean KB docx

KB_HEADER_V2 = (
    "本文件为对客安全的知识库版本，上传至售后智能体知识库后，回复生成时会全文注入"
    "作为事实依据；仅可引用本文件信息，禁止编造。"
)

KB_SECTIONS_V2: list[tuple[str, list[str]]] = [
    (
        "核心目标",
        [
            "快速解决真实问题，而不是只完成\"回复邮件\"。",
            "在不违背官网政策的前提下，降低不必要的退货、全额退款、拒付和差评。",
            "让客户感受到 LBORA 愿意负责、愿意协商、给出明确选择，而不是拖延或设置退货障碍。",
            "所有邮件都应成为可追溯记录：事实清楚、承诺清楚、下一步清楚。",
            "美国市场沟通强调直接、尊重、透明；不要把内部成本和流程压力转嫁给客户。",
        ],
    ),
    (
        "标准回复结构（7-Step）",
        [
            "称呼：Hi/Dear + 客户 First Name。",
            "确认收到：Thank you for reaching out / providing the photos.",
            "安抚：We're sorry... / We understand your frustration.",
            "事实：只写订单、照片、物流中已核实的信息。",
            "方案：最多给两个清晰方案，避免让客户继续猜下一步。",
            "行动：告诉客户现在只需要做什么，例如发照片、选 Option 1/2、提供 return tracking。",
            "落款：Best regards, + The LBORA Team。",
        ],
    ),
    (
        "语气标准",
        [
            "Warm but concise：温和但简洁，不写大段辩解。",
            "Customer-first：从客户节省时间、运费、等待的角度表达。",
            "Solution-oriented：每封售后邮件必须有明确下一步。",
            "Neutral on liability：核实前不主动承认 defective / poor quality / our fault。",
            "Consistent：与官网政策、订单、物流、退款状态一致。",
            "Professional：不争输赢、不质疑客户动机、不使用任何带指责意味的语言。",
        ],
    ),
    (
        "推荐表达与禁用表达",
        [
            "主观不满意：用 \"We're sorry the product didn't meet your expectations.\"，不用 \"Our quality is poor.\"",
            "尺码不合适：用 \"The adjustable inner band may help improve the fit.\"，不用 \"It should fit you.\"",
            "部分退款：用 \"This option may save you return shipping costs and time.\"，不用 \"Returning it is uneconomical for us.\"",
            "客户坚持退货：用 \"We completely respect your decision.\"，不用 \"You should keep it / give it to a friend.\"",
            "物流正常：用 \"Your package is still moving through the delivery network.\"，不用 \"You just need to wait.\"",
            "到货时间：用 \"We can't guarantee a specific delivery date.\"，不用 \"It will definitely arrive by Wednesday.\"",
        ],
    ),
    (
        "风险等级（决策框架）",
        [
            "P0 紧急：Chargeback / bank / credit card / consumer complaint / 严重损坏赶活动 → 停止普通挽留，优先当天处理。",
            "P1 高：第二次催退货、抱怨退货困难、明显质量照片、派送失败 → 快速给明确解决路径。",
            "P2 中：尺码、主观不喜欢、轻微变形、换货 → 可协商一次。",
            "P3 低：售前政策、库存、一般物流 → 简洁准确，促进转化。",
        ],
    ),
    (
        "总决策流程",
        [
            "确认订单号、商品、购买时间、履约/签收状态。",
            "分类：物流 / 尺码 / 主观不满意 / 客观质量 / 发错货 / 严重损坏 / 退货 / 拒付。",
            "需要证据的先要证据；物流先查 17TRACK；退款/拒付先查 Shopify 或支付渠道状态。",
            "能解决先解决：调节、恢复、补发、部分退款。",
            "首次个人原因退货可挽留一次；不要连续多轮。",
            "客户明确坚持退货或出现高风险词，立即按政策进入退货/退款。",
        ],
    ),
    (
        "尺码问题 SOP",
        [
            "核实购买尺码、商品标签、客户头围；不要只凭 \"too big\" 认定发错。",
            "有 adjustable inner band：先指导调节。",
            "正确尺码但不合适：属于 fit/personal preference，可首轮提供 Keep the hat + USD $20-30 partial refund。",
            "客户拒绝并坚持退货：停止推荐新帽子，直接按政策办理。",
            "收到错误尺码：要标签照片；确认错发后按 LBORA 责任处理，不收个人原因费用。",
        ],
    ),
    (
        "主观质量不满意 SOP",
        [
            "典型：cheaply constructed / not worth it / quality inferior / don't like the feel。",
            "没有客观证据时，不归类为 manufacturing defect。",
            "表达：didn't meet your expectations，而不是承认 defective。",
            "可一次部分退款保留商品；坚持退货则按适用政策办理。",
            "不要用\"这是热卖款\"\"别人都没问题\"反驳客户。",
        ],
    ),
    (
        "Dirty / Defective / 客观质量问题 SOP",
        [
            "必须索取：整体照片、问题部位近照、内侧标签、外包装（如相关）。",
            "收到证据前不承认 defect，不先承诺最终退款金额。",
            "轻微可处理：指导 + 合理补偿。",
            "确认制造缺陷：补发或退款，不收个人原因 $20 处理费。",
            "描述与照片不一致时，礼貌要求补充，不指责客户。",
        ],
    ),
    (
        "运输变形 / 严重损坏 SOP",
        [
            "轻微可恢复压痕：仅在适合材质时提供安全恢复建议 + 部分退款。",
            "严重断裂、开裂、帽冠帽檐分离、明显不可恢复：按 shipping damage。",
            "严重损坏：免费补发或全额退款，不收 $20。",
            "客户有明确旅行/婚礼/活动时间，补发已失去意义：优先退款。",
            "无残值严重损坏商品通常无需客户自费寄回。",
        ],
    ),
    (
        "退货 SOP：挽留阈值",
        [
            "可以挽留一次：首次提出退货；商品完好；原因是尺码/款式/个人偏好；客户语气正常，未引用政策或拒付。",
            "不要再拉扯：I just want to return / Please honor your return policy / You make it very hard to return / bank / credit card / chargeback / dispute / consumer complaint / bait and switch；或客户已明确拒绝过一次部分退款/换购方案。",
            "退货授权邮件必须写清：Your return request has been approved；只提供一个授权退货地址；商品状态要求（unworn / unused / tags / original packaging）；谁承担 return shipping；是否适用 USD $20 Return Processing Fee；寄出后提供 return tracking number；收到并检查后原路退款，时效与官网政策一致。",
        ],
    ),
    (
        "退货仓库",
        [
            "美西：Mars Shipping Service LLC, 2965 E Vernon Ave, Vernon, CA 90058, United States, +1 6264105865",
            "新泽西：Mars Shipping Service LLC, 1100 Randolph Rd. Ste B, Somerset, NJ 08873, United States, +1 6467979999",
            "规则：每封授权邮件只给一个仓库，不让客户自己选仓。",
        ],
    ),
    (
        "物流 SOP",
        [
            "先查 17TRACK / 实际承运商，不根据 Shopify 单一页面猜测。",
            "正常运输：最新节点 + 追踪方式 + 合理观察期。",
            "Flight Arrived / Customs Cleared 不等于 Delivered，也不等于丢件。",
            "Shopify waiting for details 但承运商已更新：解释同步延迟。",
            "不保证具体送达日期，除非承运商明确保证。",
            "长时间无更新：升级物流调查，不能无限期只让客户等待。",
            "派送失败：确认二次派送/自取/退回；客户已拒绝反复联系承运商时，由客服继续协调。",
        ],
    ),
    (
        "Shipping Protection SOP",
        [
            "先确认订单是否购买 Shipping Protection。",
            "丢件、损坏、符合条款的延误，按当前 SureCircle 流程处理。",
            "保险审核前不承诺\"一定赔付\"。",
            "商家决定直接退款/补发前，核对保险理赔状态，避免重复赔付。",
        ],
    ),
    (
        "Chargeback / Dispute SOP",
        [
            "出现 bank / credit card / chargeback / dispute：立即按最高优先级处理。",
            "优先尝试直接解决，但不做与支付渠道规则冲突的退款承诺。",
            "正式 Chargeback 开始后先查看 Shopify/支付渠道状态；不能退款时不要承诺\"马上原路退款\"。",
            "即使客户称已撤销，也要按支付渠道要求提交撤销证明/证据，直到后台真正结束。",
        ],
    ),
    (
        "售前邮件标准",
        [
            "售前回答简洁、促进决策，不主动把退货费用写成一大段警告。",
            "客户问 Return Policy：说明 30-day return、商品状态要求、个人原因成本、质量/错发由品牌负责。",
            "客户问尺码：先问头围，再推荐尺码；适用款提醒 adjustable inner band。",
            "客户问能否赶日期：说明处理和运输逻辑，不保证不可控日期。",
            "客户问 Gift Card：先确认网站当前是否有礼品卡产品，再回复；不能凭印象说有或没有。",
        ],
    ),
    (
        "发送前检查（对客质量清单）",
        [
            "客户名字是否正确？订单号/SKU 是否正确？",
            "是否查过订单状态？物流是否查过最新节点？",
            "客户说的是主观不满意还是客观缺陷？需要照片是否已经索取？",
            "有没有未经核实承认责任？",
            "是否给了清晰方案和下一步？",
            "客户已经坚持退货时是否还在拉扯？",
            "金额、$20、退货运费是否符合当前政策？是否只给一个正确退货仓？",
            "落款是否统一为 The LBORA Team？",
        ],
    ),
    (
        "标准邮件骨架",
        [
            "Dear [Name],",
            "Thank you for reaching out / providing the photos.",
            "We're sorry that [the fit / delivery / product] did not meet your expectations, and we understand your concern.",
            "We have reviewed [verified fact].",
            "Option 1: [solution]",
            "Option 2: [solution, if needed]",
            "Please [single clear next action].",
            "Best regards,",
            "The LBORA Team",
        ],
    ),
]


def _add_heading(doc: docx.Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def _add_bullets(doc: docx.Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def _add_block(doc: docx.Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)


def build_kb_docx(out_path: Path, cases: list[dict]) -> None:
    doc = docx.Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("LBORA 美国市场客户邮件回复标准 SOP（纯净版）")
    tr.bold = True
    tr.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Customer Email Response Standards & Resolution Playbook · V2.0")
    sr.font.size = Pt(12)

    note = doc.add_paragraph()
    note.add_run(KB_HEADER_V2).font.size = Pt(10)
    note.paragraph_format.space_after = Pt(8)

    for i, (heading, items) in enumerate(KB_SECTIONS_V2, start=1):
        _add_heading(doc, f"{i}. {heading}")
        _add_bullets(doc, items)

    if cases:
        _add_heading(doc, f"{len(KB_SECTIONS_V2) + 1}. 真实案例参考（已发出的标准英文话术）")
        for c in cases:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(f"案例{c['num']:02d}｜{c['title']}")
            r.bold = True
            r.font.size = Pt(11)
            _add_block(doc, c["reply"].split("\n"))

    doc.save(str(out_path))


# ----------------------------------------------------------------------- main

def main() -> None:
    parser = argparse.ArgumentParser(description="Build LBORA importable QA + clean KB assets")
    parser.add_argument("--faq", type=Path, default=REPO_ROOT / "LBORA_客服标准FAQ_V1.0.docx")
    parser.add_argument("--cases", type=Path, default=REPO_ROOT / "LBORA_7月7日起客户沟通案例与客服标准_V1.0 (1).docx")
    parser.add_argument("--out-json", type=Path, default=REPO_ROOT / "qa_pairs_import.json")
    parser.add_argument("--out-docx", type=Path, default=REPO_ROOT / "LBORA_客服知识库_纯净版_V2.0.docx")
    args = parser.parse_args()

    if not args.faq.exists():
        raise SystemExit(f"source FAQ not found: {args.faq}")

    entries = parse_faq(args.faq)

    cases: list[dict] = []
    if args.cases.exists():
        cases = parse_cases(args.cases)
        by_num = {c["num"]: c for c in cases}
        missing = [num for num, _, _ in CASE_QA if num not in by_num]
        if missing:
            raise SystemExit(f"case study docx missing expected cases: {missing}")
        for num, question, category in CASE_QA:
            entries.append({"question": question, "answer": by_num[num]["reply"], "category": category})
    print(f"[ok] FAQ QA: {len(entries) - len(CASE_QA)}; case QA: +{len(CASE_QA)}")

    payload = {"items": entries}
    args.out_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[ok] total QA entries: {len(entries)} -> {args.out_json}")

    build_kb_docx(args.out_docx, cases)
    print(f"[ok] clean KB V2.0 docx (with {len(cases)} case examples) -> {args.out_docx}")


if __name__ == "__main__":
    main()

"""Knowledge base service (M-08).

Upload PDF/DOCX/MD, extract plain text and store it in full (no chunking,
no vectorization, no retrieval). Reply generation injects the full text
directly into the prompt (TECH M-07 / M-08 red line).
"""

from __future__ import annotations

import io
import logging
import os

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.exceptions import (
    KnowledgeEmptyError,
    KnowledgeError,
    KnowledgeUnsupportedError,
)
from app.models.knowledge_doc import KnowledgeDoc
from app.services.audit import utcnow

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md"}
MAX_KB_BYTES = 20 * 1024 * 1024


def _extract_pdf(data: bytes) -> str:
    """Extract all page text from a PDF (pypdf, pure text extraction)."""

    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency declared
        raise KnowledgeError("pypdf is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001 - encrypted pdfs unsupported
                raise KnowledgeEmptyError("PDF is encrypted and cannot be read") from exc
        pages = [page.extract_text() or "" for page in reader.pages]
    except KnowledgeEmptyError:
        raise
    except Exception as exc:  # noqa: BLE001 - malformed pdfs raise many types
        raise KnowledgeError(f"Failed to parse PDF: {exc}") from exc
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    """Extract paragraphs + table cells from a DOCX (python-docx)."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency declared
        raise KnowledgeError("python-docx is not installed") from exc
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:  # noqa: BLE001 - malformed docx raises many types
        raise KnowledgeError(f"Failed to parse DOCX: {exc}") from exc
    return "\n".join(parts).strip()


def extract_text(filename: str, data: bytes) -> str:
    """Return the plain text of a pdf/docx/md upload."""

    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".md":
        return data.decode("utf-8", errors="replace")
    raise KnowledgeUnsupportedError(
        f"Unsupported file type {ext or '(none)'}; only pdf/docx/md are allowed"
    )


class KnowledgeService:
    """CRUD + full-text loading for knowledge documents."""

    def __init__(self, db: Session) -> None:
        self.db = db

    def upload(self, filename: str, data: bytes) -> KnowledgeDoc:
        """Insert a new doc or bump the version of the active doc with the
        same filename (overwrite semantics, TECH M-08)."""

        if len(data) > MAX_KB_BYTES:
            raise KnowledgeError("Upload exceeds the 20MB limit")
        safe_name = os.path.basename(filename).replace("/", "_").replace("\\", "_")
        safe_name = safe_name or "document"
        text = extract_text(safe_name, data).strip()
        if not text:
            raise KnowledgeEmptyError("No extractable text found in the file")

        existing = self.db.execute(
            select(KnowledgeDoc).where(
                KnowledgeDoc.filename == safe_name,
                KnowledgeDoc.is_deleted.is_(False),
            )
        ).scalars().first()
        if existing is not None:
            existing.version += 1
            existing.content = text
            existing.uploaded_at = utcnow()
            doc = existing
            logger.info("Knowledge doc %r overwritten -> v%s", safe_name, existing.version)
        else:
            doc = KnowledgeDoc(
                filename=safe_name,
                version=1,
                content=text,
                uploaded_at=utcnow(),
            )
            self.db.add(doc)
            logger.info("Knowledge doc %r uploaded v1", safe_name)
        self.db.flush()
        return doc

    def list_docs(self) -> list[KnowledgeDoc]:
        return self.db.execute(
            select(KnowledgeDoc)
            .where(KnowledgeDoc.is_deleted.is_(False))
            .order_by(KnowledgeDoc.uploaded_at.desc(), KnowledgeDoc.id.desc())
        ).scalars().all()

    def get(self, doc_id: int) -> KnowledgeDoc | None:
        doc = self.db.get(KnowledgeDoc, doc_id)
        if doc is None or doc.is_deleted:
            return None
        return doc

    def soft_delete(self, doc_id: int) -> bool:
        doc = self.db.get(KnowledgeDoc, doc_id)
        if doc is None or doc.is_deleted:
            return False
        doc.is_deleted = True
        self.db.flush()
        return True

    def full_text(self) -> str:
        """All active docs joined for full injection into the reply prompt."""

        docs = self.db.execute(
            select(KnowledgeDoc)
            .where(KnowledgeDoc.is_deleted.is_(False))
            .order_by(KnowledgeDoc.id.asc())
        ).scalars().all()
        return "\n\n".join(
            f"[{doc.filename} v{doc.version}]\n{doc.content}" for doc in docs
        )
